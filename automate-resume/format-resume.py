import re
import os
import argparse

# Library imports for DOCX and PDF generation
import docx
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, KeepTogether
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

def parse_txt_resume(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Splits sections cleanly even with symbols like '&'
    sections = re.split(r'\n(?=[A-Z\s&/]{4,}\n)', content)
    
    header_raw = sections[0].strip().split('\n')
    name = header_raw[0]
    title = header_raw[1] if len(header_raw) > 1 else ""
    contact_info = " | ".join(header_raw[2:]) if len(header_raw) > 2 else ""

    parsed_sections = {}
    for section in sections[1:]:
        lines = section.strip().split('\n')
        if not lines:
            continue
        section_title = lines[0].strip()
        section_body = "\n".join(lines[1:]).strip()
        parsed_sections[section_title] = section_body

    return {
        "name": name,
        "title": title,
        "contact": contact_info,
        "sections": parsed_sections
    }

def generate_html_resume(data, output_path):
    css = """
    @page { size: A4; margin: 20mm 15mm; }
    body { font-family: 'Segoe UI', Helvetica, Arial, sans-serif; color: #1e293b; margin: 0; padding: 0; font-size: 10pt; line-height: 1.5; }
    .header { text-align: center; margin-bottom: 22px; border-bottom: 2px solid #cbd5e1; padding-bottom: 12px; }
    h1 { font-size: 22pt; color: #0f172a; margin: 0 0 4px 0; text-transform: uppercase; font-weight: 700; letter-spacing: 0.5px; }
    .subtitle { font-size: 11pt; color: #2563eb; font-weight: 600; margin: 0 0 8px 0; }
    .contact { font-size: 9pt; color: #64748b; }
    h2 { font-size: 12pt; color: #0f172a; border-left: 3.5px solid #2563eb; padding-left: 8px; margin: 18px 0 10px 0; text-transform: uppercase; font-weight: 700; page-break-after: avoid; }
    p { margin: 0 0 10px 0; text-align: justify; font-weight: normal; }
    ul { margin: 0 0 10px 0; padding-left: 20px; }
    li { margin-bottom: 4px; text-align: justify; }
    .experience-block { margin-bottom: 16px; page-break-inside: avoid; }
    .exp-header { font-weight: bold; color: #0f172a; display: flex; justify-content: space-between; margin-bottom: 2px; }
    .exp-scope { font-style: italic; color: #475569; margin-bottom: 4px; font-size: 9.5pt; }
    .exp-stack { font-size: 8.5pt; color: #1e293b; margin-top: 4px; background-color: #f1f5f9; padding: 3px 6px; border-radius: 3px; display: inline-block; }
    .exp-stack strong { color: #475569; }
    """

    html_content = f"""<!DOCTYPE html><html><head><meta charset="utf-8"><title>{data['name']}</title><style>{css}</style></head><body>
    <div class="header"><h1>{data['name']}</h1><div class="subtitle">{data['title']}</div><div class="contact">{data['contact']}</div></div>"""

    for title, body in data['sections'].items():
        html_content += f"    <h2>{title.title()}</h2>\n"
        if "EXPERIENCE" in title and "EARLIER" not in title:
            jobs = re.split(r'\n(?=[A-Za-z0-9\s/&]+?—\s)', '\n' + body)
            for job in jobs:
                if not job.strip(): continue
                job_lines = job.strip().split('\n')
                meta_line, scope, stack, bullets = job_lines[0], "", "", []
                for line in job_lines[1:]:
                    if line.strip().startswith("Scope:"): scope = line.strip()
                    elif line.strip().startswith("Stack:"): stack = line.strip()
                    elif line.strip().startswith("•") or line.strip().startswith("-"): bullets.append(line.strip()[1:].strip())
                    elif line.strip(): bullets.append(line.strip())

                if " | " in meta_line:
                    left_side, right_side = meta_line.split(" | ", 1)
                    header_html = f'<div class="exp-header"><span>{left_side}</span><span style="font-size: 9.5pt; color: #64748b; font-weight: normal;">{right_side}</span></div>'
                else:
                    header_html = f'<div class="exp-header"><span>{meta_line}</span></div>'

                html_content += f'    <div class="experience-block">\n        {header_html}\n'
                if scope: html_content += f'        <div class="exp-scope">{scope}</div>\n'
                if bullets:
                    html_content += '        <ul>\n'
                    for b in bullets: html_content += f'            <li>{b}</li>\n'
                    html_content += '        </ul>\n'
                if stack:
                    html_content += f'        <div class="exp-stack">{stack.replace("Stack:", "<strong>Stack:</strong>")}</div>\n'
                html_content += '    </div>\n'
        else:
            lines = body.split('\n')
            is_list = any(l.strip().startswith('•') or l.strip().startswith('-') for l in lines) or "EXPERIENCE" in title
            if is_list:
                html_content += "    <ul style='list-style-type: none; padding-left: 5px;'>\n"
                for line in lines:
                    if line.strip():
                        clean_line = line.strip().lstrip('•-').strip()
                        colon_index = clean_line.find(":")
                        if colon_index > 0 and not clean_line.startswith("http"):
                            k, v = clean_line[:colon_index].strip(), clean_line[colon_index+1:]
                            if len(k) < 50 and "—" not in clean_line or any(w in k.lower() for w in ["bachelor", "degree", "university", "universidad", "analista"]):
                                clean_line = f"<strong>{k}:</strong>{v}"
                            else:
                                clean_line = f"<span style='font-weight: normal;'>{clean_line}</span>"
                        else:
                            clean_line = f"<span style='font-weight: normal;'>{clean_line}</span>"
                        html_content += f"        <li style='margin-bottom: 5px;'>{clean_line}</li>\n"
                html_content += "    </ul>\n"
            else:
                for p in body.split('\n\n'):
                    if p.strip(): html_content += f"    <p>{p.strip()}</p>\n"

    html_content += "</body></html>"
    with open(output_path, 'w', encoding='utf-8') as f: f.write(html_content)


def generate_docx_resume(data, output_path):
    doc = Document()
    # Set Margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Base Styles Configuration
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Arial'
    style_normal.font.size = Pt(10)
    style_normal.font.color.rgb = docx.shared.RGBColor(30, 41, 59)

    # Header
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_name = p_name.add_run(data['name'].upper())
    run_name.font.size = Pt(22)
    run_name.font.bold = True
    run_name.font.color.rgb = docx.shared.RGBColor(15, 23, 42)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run(data['title'])
    run_title.font.size = Pt(11)
    run_title.font.bold = True
    run_title.font.color.rgb = docx.shared.RGBColor(37, 99, 235)

    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_after = Pt(18)
    run_contact = p_contact.add_run(data['contact'])
    run_contact.font.size = Pt(9)
    run_contact.font.color.rgb = docx.shared.RGBColor(100, 116, 139)

    # Add a thin border line
    p_border = doc.add_paragraph()
    p_border.paragraph_format.space_after = Pt(12)
    pBrd = parse_xml(f'<w:pBrd {nsdecls("w")}><w:bottom w:val="single" w:sz="12" w:space="1" w:color="CBD5E1"/></w:pBrd>')
    p_border._p.get_or_add_pPr().append(pBrd)

    for title, body in data['sections'].items():
        # Section Header
        h2 = doc.add_paragraph()
        h2.paragraph_format.space_before = Pt(14)
        h2.paragraph_format.space_after = Pt(6)
        h2.paragraph_format.keep_with_next = True
        
        # Accent left border via XML
        pBrd_left = parse_xml(f'<w:pBrd {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="4" w:color="2563EB"/></w:pBrd>')
        h2._p.get_or_add_pPr().append(pBrd_left)
        
        run_h2 = h2.add_run(f"  {title.upper()}")
        run_h2.font.size = Pt(12)
        run_h2.font.bold = True
        run_h2.font.color.rgb = docx.shared.RGBColor(15, 23, 42)

        if "EXPERIENCE" in title and "EARLIER" not in title:
            jobs = re.split(r'\n(?=[A-Za-z0-9\s/&]+?—\s)', '\n' + body)
            for job in jobs:
                if not job.strip(): continue
                job_lines = job.strip().split('\n')
                meta_line, scope, stack, bullets = job_lines[0], "", "", []
                for line in job_lines[1:]:
                    if line.strip().startswith("Scope:"): scope = line.strip()
                    elif line.strip().startswith("Stack:"): stack = line.strip()
                    elif line.strip().startswith("•") or line.strip().startswith("-"): bullets.append(line.strip()[1:].strip())
                    elif line.strip(): bullets.append(line.strip())

                # Top line of experience
                p_job = doc.add_paragraph()
                p_job.paragraph_format.space_after = Pt(2)
                p_job.paragraph_format.keep_with_next = True
                
                if " | " in meta_line:
                    left, right = meta_line.split(" | ", 1)
                    r_left = p_job.add_run(left)
                    r_left.font.bold = True
                    r_space = p_job.add_run(" " * (90 - len(meta_line))) # basic tab stop approximation
                    r_right = p_job.add_run(f" | {right}")
                    r_right.font.color.rgb = docx.shared.RGBColor(100, 116, 139)
                else:
                    r_meta = p_job.add_run(meta_line)
                    r_meta.font.bold = True

                if scope:
                    p_scope = doc.add_paragraph()
                    p_scope.paragraph_format.space_after = Pt(3)
                    r_sc = p_scope.add_run(scope)
                    r_sc.font.italic = True
                    r_sc.font.color.rgb = docx.shared.RGBColor(71, 85, 110)

                for b in bullets:
                    p_b = doc.add_paragraph(style='List Bullet')
                    p_b.paragraph_format.space_after = Pt(2)
                    p_b.add_run(b)

                if stack:
                    p_st = doc.add_paragraph()
                    p_st.paragraph_format.space_before = Pt(2)
                    p_st.paragraph_format.space_after = Pt(8)
                    p_st.add_run("Stack: ").font.bold = True
                    p_st.add_run(stack.replace("Stack:", "").strip())

        else:
            lines = body.split('\n')
            is_list = any(l.strip().startswith('•') or l.strip().startswith('-') for l in lines) or "EXPERIENCE" in title
            if is_list:
                for line in lines:
                    if not line.strip(): continue
                    clean_line = line.strip().lstrip('•-').strip()
                    p_l = doc.add_paragraph()
                    p_l.paragraph_format.space_after = Pt(3)
                    
                    colon_index = clean_line.find(":")
                    if colon_index > 0 and "—" not in clean_line:
                        k = clean_line[:colon_index].strip()
                        v = clean_line[colon_index+1:]
                        k_lower = k.lower()
                        if len(k) < 50 or any(w in k_lower for w in ["bachelor", "degree", "university", "universidad", "analista"]):
                            p_l.add_run(f"• {k}: ").font.bold = True
                            p_l.add_run(v)
                            continue
                    
                    p_l.add_run(f"• {clean_line}")
            else:
                for p in body.split('\n\n'):
                    if p.strip():
                        doc.add_paragraph(p.strip()).paragraph_format.space_after = Pt(8)

    doc.save(output_path)


def generate_pdf_resume(data, output_path):
    doc = SimpleDocTemplate(output_path, pagesize=letter, leftMargin=40, rightMargin=40, topMargin=40, bottomMargin=40)
    styles = getSampleStyleSheet()
    
    # Custom PDF Paragraph Styles
    style_name = ParagraphStyle('Name', fontName='Helvetica-Bold', fontSize=22, leading=26, textColor=colors.HexColor('#0f172a'), alignment=1)
    style_title = ParagraphStyle('Title', fontName='Helvetica-Bold', fontSize=11, leading=15, textColor=colors.HexColor('#2563eb'), alignment=1, spaceBefore=4)
    style_contact = ParagraphStyle('Contact', fontName='Helvetica', fontSize=9, leading=13, textColor=colors.HexColor('#64748b'), alignment=1, spaceBefore=4, spaceAfter=12)
    
    style_h2 = ParagraphStyle('H2', fontName='Helvetica-Bold', fontSize=12, leading=16, textColor=colors.HexColor('#0f172a'), spaceBefore=14, spaceAfter=6, keepWithNext=True)
    style_body = ParagraphStyle('Body', fontName='Helvetica', fontSize=9.5, leading=14, textColor=colors.HexColor('#1e293b'), alignment=4, spaceAfter=8)
    style_bullet = ParagraphStyle('Bullet', fontName='Helvetica', fontSize=9.5, leading=14, textColor=colors.HexColor('#1e293b'), leftIndent=15, firstLineIndent=-10, spaceAfter=3)
    
    style_job_header = ParagraphStyle('JobHeader', fontName='Helvetica-Bold', fontSize=10, leading=14, textColor=colors.HexColor('#0f172a'), keepWithNext=True)
    style_job_scope = ParagraphStyle('JobScope', fontName='Helvetica-Oblique', fontSize=9.5, leading=13, textColor=colors.HexColor('#475569'), spaceAfter=4, keepWithNext=True)
    style_job_stack = ParagraphStyle('JobStack', fontName='Helvetica', fontSize=8.5, leading=12, textColor=colors.HexColor('#1e293b'), spaceBefore=3, spaceAfter=10)

    story = []
    
    # Build Header
    story.append(Paragraph(data['name'].upper(), style_name))
    story.append(Paragraph(data['title'], style_title))
    story.append(Paragraph(data['contact'], style_contact))
    
    # Custom Horizontal Divider Rule
    story.append(Spacer(1, 2))
    
    for title, body in data['sections'].items():
        # Section Titles
        story.append(Paragraph(f"<font color='#2563eb'>|</font> &nbsp;{title.upper()}", style_h2))
        
        if "EXPERIENCE" in title and "EARLIER" not in title:
            jobs = re.split(r'\n(?=[A-Za-z0-9\s/&]+?—\s)', '\n' + body)
            for job in jobs:
                if not job.strip(): continue
                job_elements = []
                job_lines = job.strip().split('\n')
                meta_line, scope, stack, bullets = job_lines[0], "", "", []
                for line in job_lines[1:]:
                    if line.strip().startswith("Scope:"): scope = line.strip()
                    elif line.strip().startswith("Stack:"): stack = line.strip()
                    elif line.strip().startswith("•") or line.strip().startswith("-"): bullets.append(line.strip()[1:].strip())
                    elif line.strip(): bullets.append(line.strip())

                if " | " in meta_line:
                    left, right = meta_line.split(" | ", 1)
                    meta_formatted = f"{left} <font color='#64748b' size='9'>| {right}</font>"
                else:
                    meta_formatted = meta_line
                
                job_elements.append(Paragraph(meta_formatted, style_job_header))
                if scope: job_elements.append(Paragraph(scope, style_job_scope))
                for b in bullets:
                    job_elements.append(Paragraph(f"&bull; {b}", style_bullet))
                if stack:
                    stack_formatted = stack.replace("Stack:", "<b>Stack:</b>")
                    job_elements.append(Paragraph(stack_formatted, style_job_stack))
                
                story.append(KeepTogether(job_elements))
        else:
            lines = body.split('\n')
            is_list = any(l.strip().startswith('•') or l.strip().startswith('-') for l in lines) or "EXPERIENCE" in title
            if is_list:
                for line in lines:
                    if not line.strip(): continue
                    clean_line = line.strip().lstrip('•-').strip()
                    colon_index = clean_line.find(":")
                    if colon_index > 0 and "—" not in clean_line:
                        k, v = clean_line[:colon_index].strip(), clean_line[colon_index+1:]
                        k_lower = k.lower()
                        if len(k) < 50 or any(w in k_lower for w in ["bachelor", "degree", "university", "universidad", "analista"]):
                            clean_line = f"<b>{k}:</b>{v}"
                    story.append(Paragraph(f"&bull; {clean_line}", style_bullet))
            else:
                for p in body.split('\n\n'):
                    if p.strip(): story.append(Paragraph(p.strip(), style_body))

    doc.build(story)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Convert plain text resumes into beautiful HTML, DOCX, and PDF formats.")
    parser.add_argument("-i", "--input", required=True, help="Path to input .txt file")
    parser.add_argument("-o", "--output", required=False, help="Base path for output files (without extension)")
    args = parser.parse_args()

    if not os.path.exists(args.input):
        raise FileNotFoundError(f"Input file '{args.input}' not found.")

    base_path = args.output
    if not base_path:
        base_path, _ = os.path.splitext(args.input)
        base_path = f"{base_path}_out"

    # Define outputs paths
    html_out = f"{base_path}.html"
    docx_out = f"{base_path}.docx"
    pdf_out = f"{base_path}.pdf"

    # Run structural extraction engines
    resume_data = parse_txt_resume(args.input)
    
    generate_html_resume(resume_data, html_out)
    generate_docx_resume(resume_data, docx_out)
    generate_pdf_resume(resume_data, pdf_out)
    
    print(f"\n🚀 Complete compilation successful!\n📄 HTML: {html_out}\n📝 DOCX: {docx_out}\n🛡️ PDF:  {pdf_out}")